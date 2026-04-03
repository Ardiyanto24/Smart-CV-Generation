# cv-agent/backend/renderer/docx_renderer.py

"""
DOCX Renderer — Document Renderer Layer

Mengkonversi Final Structured Output JSON menjadi DOCX bytes
menggunakan python-docx dan named styles dari cv_template.docx.

Satu fungsi publik:
- render_docx(cv_output) -> bytes
"""

import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm

logger = logging.getLogger("renderer.docx_renderer")

# ── Template path ─────────────────────────────────────────────────────────────
# Path absolut dihitung dari lokasi file ini — konsisten dengan pdf_renderer.py
_TEMPLATE_PATH = Path(__file__).parent / "templates" / "cv_template.docx"


def _add_right_aligned_date(paragraph, date_text: str):
    """
    Tambahkan teks tanggal yang right-aligned ke paragraf yang sudah ada
    menggunakan tab stop di right margin.

    Word menggunakan tab stop untuk right-align teks dalam satu paragraf.
    Kita inject tab stop via OOXML karena python-docx tidak expose
    paragraph tab stops via API tingkat tinggi.
    """
    # Set tab stop di right margin (16cm untuk A4 dengan margin 15mm kiri-kanan)
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9070")   # 16cm dalam twips (1cm = 567 twips)
    tabs.append(tab)
    pPr.append(tabs)

    # Tambahkan tab character + teks tanggal ke paragraf
    # Run pertama sudah ada (entry title) — kita tambahkan run baru
    run = paragraph.add_run(f"\t{date_text}")
    run.font.size = Pt(9)
    run.font.bold = False


def _add_section_heading(doc: Document, text: str):
    """Helper: tambahkan section heading dengan style CV Section Heading."""
    doc.add_paragraph(text.upper(), style="CV Section Heading")


def _build_contact_line(header: dict) -> str:
    """
    Gabungkan semua contact fields yang tidak kosong dengan separator ' | '.
    Field yang None atau empty string di-skip.
    """
    fields = [
        header.get("email"),
        header.get("phone"),
        header.get("linkedin"),
        header.get("github"),
        header.get("portfolio"),
    ]
    return " | ".join(f for f in fields if f)


def render_docx(cv_output: dict) -> bytes:
    """
    Render Final Structured Output JSON menjadi DOCX bytes.

    Args:
        cv_output: Final Structured Output dict dari Cluster 5.
                   Section yang kosong atau tidak ada di-skip tanpa error.

    Returns:
        Raw DOCX bytes — siap di-upload ke Supabase Storage.

    Raises:
        FileNotFoundError: jika cv_template.docx tidak ditemukan
        Exception        : python-docx errors dibiarkan propagate ke orchestrator
    """
    logger.info("[render_docx] starting DOCX render")

    # ── Load template ─────────────────────────────────────────────────────────
    if not _TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Template 'cv_template.docx' tidak ditemukan di: {_TEMPLATE_PATH}. "
            f"Jalankan create_template.py terlebih dahulu."
        )

    doc = Document(str(_TEMPLATE_PATH))

    # Hapus semua paragraph default yang ada di template kosong
    # Template baru dari python-docx selalu punya satu paragraph kosong
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # ── 1. HEADER ─────────────────────────────────────────────────────────────
    header = cv_output.get("header", {})
    if header and header.get("name"):
        doc.add_paragraph(header["name"], style="CV Name")

        contact_line = _build_contact_line(header)
        if contact_line:
            doc.add_paragraph(contact_line, style="CV Contact")

    # ── 2. SUMMARY ────────────────────────────────────────────────────────────
    summary = cv_output.get("summary")
    if summary:
        _add_section_heading(doc, "Summary")
        doc.add_paragraph(summary, style="CV Summary")

    # ── 3. EXPERIENCE ─────────────────────────────────────────────────────────
    experience = cv_output.get("experience", [])
    if experience:
        _add_section_heading(doc, "Experience")
        for entry in experience:
            # Entry title + date dalam satu paragraf
            p = doc.add_paragraph(
                entry.get("company", ""),
                style="CV Entry Title",
            )
            if entry.get("year"):
                _add_right_aligned_date(p, entry["year"])

            # Subtitle: role
            if entry.get("role"):
                doc.add_paragraph(entry["role"], style="CV Entry Subtitle")

            # Bullets
            for bullet in entry.get("bullets", []):
                doc.add_paragraph(f"• {bullet}", style="CV Bullet")

    # ── 4. EDUCATION ──────────────────────────────────────────────────────────
    education = cv_output.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for entry in education:
            # Entry title: institution + date
            p = doc.add_paragraph(
                entry.get("institution", ""),
                style="CV Entry Title",
            )
            if entry.get("year"):
                _add_right_aligned_date(p, entry["year"])

            # Subtitle: degree + field + GPA
            subtitle_parts = [entry.get("degree", "")]
            if entry.get("field"):
                subtitle_parts.append(entry["field"])
            if entry.get("gpa"):
                subtitle_parts.append(f"GPA {entry['gpa']}")
            subtitle = " — ".join(p for p in subtitle_parts if p)
            if subtitle:
                doc.add_paragraph(subtitle, style="CV Entry Subtitle")

            # Bullets
            for bullet in entry.get("bullets", []):
                doc.add_paragraph(f"• {bullet}", style="CV Bullet")

    # ── 5. AWARDS ─────────────────────────────────────────────────────────────
    awards = cv_output.get("awards", [])
    if awards:
        _add_section_heading(doc, "Awards")
        for entry in awards:
            p = doc.add_paragraph(
                entry.get("title", ""),
                style="CV Entry Title",
            )
            if entry.get("year"):
                _add_right_aligned_date(p, entry["year"])

            if entry.get("issuer"):
                doc.add_paragraph(entry["issuer"], style="CV Entry Subtitle")

            for bullet in entry.get("bullets", []):
                doc.add_paragraph(f"• {bullet}", style="CV Bullet")

    # ── 6. SKILLS ─────────────────────────────────────────────────────────────
    skills = cv_output.get("skills", {})
    skills_grouped = skills.get("skills_grouped", []) if skills else []
    if skills_grouped:
        _add_section_heading(doc, "Skills")
        for group in skills_grouped:
            label = group.get("group_label", "")
            items = ", ".join(group.get("items", []))
            p = doc.add_paragraph(style="CV Skills Group")
            # Label bold, items normal — dua run dalam satu paragraf
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            p.add_run(items)

    # ── 7. PROJECTS ───────────────────────────────────────────────────────────
    projects = cv_output.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for entry in projects:
            p = doc.add_paragraph(
                entry.get("title", ""),
                style="CV Entry Title",
            )
            # Date kolom kanan: github_url jika ada
            if entry.get("github_url"):
                _add_right_aligned_date(p, entry["github_url"])

            # Subtitle: tools yang dipakai
            tools = entry.get("tools", [])
            if tools:
                doc.add_paragraph(", ".join(tools), style="CV Entry Subtitle")

            for bullet in entry.get("bullets", []):
                doc.add_paragraph(f"• {bullet}", style="CV Bullet")

    # ── 8. CERTIFICATES ───────────────────────────────────────────────────────
    certificates = cv_output.get("certificates", [])
    if certificates:
        _add_section_heading(doc, "Certificates")
        for cert in certificates:
            # Certificates adalah pass-through — tidak ada bullets
            # Format: "Nama Sertifikat · Issuer · Year"
            parts = [cert.get("name", "")]
            if cert.get("issuer"):
                parts.append(cert["issuer"])
            if cert.get("year"):
                parts.append(str(cert["year"]))
            line = " · ".join(p for p in parts if p)
            doc.add_paragraph(line, style="CV Skills Group")

    # ── 9. ORGANIZATIONS ──────────────────────────────────────────────────────
    organizations = cv_output.get("organizations", [])
    if organizations:
        _add_section_heading(doc, "Organizations")
        for entry in organizations:
            p = doc.add_paragraph(
                entry.get("name", ""),
                style="CV Entry Title",
            )
            if entry.get("year"):
                _add_right_aligned_date(p, entry["year"])

            if entry.get("role"):
                doc.add_paragraph(entry["role"], style="CV Entry Subtitle")

            for bullet in entry.get("bullets", []):
                doc.add_paragraph(f"• {bullet}", style="CV Bullet")

    # ── Simpan ke BytesIO ─────────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()

    logger.info(
        f"[render_docx] DOCX render complete, size={len(docx_bytes)} bytes"
    )

    return docx_bytes
